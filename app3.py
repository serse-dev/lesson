import streamlit as st
import pandas as pd
import os
import re
import google.generativeai as genai
from docx import Document
from io import BytesIO

# --- Gemini API Key (шууд код дотор бичсэн) ---
genai.configure(api_key="AIzaSyCB2-tcGlXPEgbc9_jWRH4vwltkm7t15a0")
model = genai.GenerativeModel("gemini-1.5-flash")

# --- Файлын зам ---
csv_path = "csv_exports/Plan.csv"
shalg_path = "csv_exports/Шалгуур.csv"
criteria_path = "csv_exports/Үр дүнгийн шалгуур.csv"
level_path = "csv_exports/Гүйцэтгэлийн түвшин.csv"

# --- Streamlit тохиргоо ---
st.set_page_config(page_title="Ээлжит хичээлийн төлөвлөлт", page_icon="📚", layout="wide")
st.title("📚 Ээлжит хичээлийн төлөвлөлт")

def split_sentences(text):
    if not isinstance(text, str):
        return []
    sentences = [s.strip() for s in re.split(r"\.(?!\d)", text) if s.strip()]
    sentences = [s + "." if not s.endswith(".") else s for s in sentences]
    filtered = [s for s in sentences if not re.match(r"^\d+\.\d+\.*$", s)]
    return filtered

# --- Гол хэсэг ---
if not os.path.exists(csv_path):
    st.error(f"{csv_path} файл олдсонгүй. Upload хэсгээр оруулна уу.")
    uploaded = st.file_uploader("CSV файл оруулах", type=["csv"])
    if uploaded:
        df = pd.read_csv(uploaded)
else:
    df = pd.read_csv(csv_path)

if "df" in locals():
    df = df.rename(columns={
        "Судлагдахууны нэр": "Хичээлийн нэр",
        "Нэгжийн нэр": "Нэгж хичээл",
        "Ээлжит хичээл": "Ээлжит хичээл",
        "Ээлжит хичээлийн зорилго": "Ээлжит хичээлийн зорилго"
    })

    subjects = df["Хичээлийн нэр"].unique()
    subject = st.selectbox("Хичээлийн нэр сонгох", subjects)
    unit_df = df[df["Хичээлийн нэр"] == subject]
    units = unit_df["Нэгж хичээл"].unique()
    unit = st.selectbox("Нэгж хичээл сонгох", units)
    lesson_df = unit_df[unit_df["Нэгж хичээл"] == unit]
    lessons = lesson_df["Ээлжит хичээл"].unique()
    lesson = st.selectbox("Ээлжит хичээл сонгох", lessons)
    selected = lesson_df[lesson_df["Ээлжит хичээл"] == lesson]

    if selected.empty:
        st.warning("Тохирох ээлжит хичээл олдсонгүй.")
    else:
        selected = selected.iloc[0]

        # --- Суралцахуйн зорилт ---
        learning_objective = "Хоосон байна."
        learning_outcome = "Хоосон байна."
        objectives_list = []
        outcomes_list = []
        if os.path.exists(shalg_path):
            shalg_df = pd.read_csv(shalg_path)
            shalg_match = shalg_df[shalg_df["Ээлжит хичээл"] == lesson]
            if not shalg_match.empty:
                objectives_list = shalg_match["Суралцахуйн зорилт"].dropna().unique().tolist()
                outcomes_list = shalg_match["Суралцахуйн үр дүн"].dropna().unique().tolist()
                learning_objective = "\n".join(objectives_list) if objectives_list else "Хоосон байна."
                learning_outcome = "\n".join(outcomes_list) if outcomes_list else "Хоосон байна."

        # --- Мэдээлэл харуулах ---
        st.markdown(f"""
        **Хичээлийн нэр:** {selected['Хичээлийн нэр']}  
        **Нэгж хичээл:** {selected['Нэгж хичээл']}  
        **Ээлжит хичээл:** {selected['Ээлжит хичээл']}  
        **Суралцахуйн зорилт:** {learning_objective}  
        **Ээлжит хичээлийн зорилго:** {selected['Ээлжит хичээлийн зорилго']}  
        **Суралцахуйн үр дүн:** {learning_outcome}  
        """)

        # --- Prompt үүсгэх ---
        st.markdown("#### 📋 Prompt үүсгэх")

        default_prompt = f"""
Та дараах мэдээлэл дээр тулгуурлан ээлжит хичээлийн дэлгэрэнгүй төлөвлөлт гарга:
- Хичээлийн нэр: {selected['Хичээлийн нэр']}
- Нэгж хичээл: {selected['Нэгж хичээл']}
- Ээлжит хичээл: {selected['Ээлжит хичээл']}
- Суралцахуйн зорилт: {learning_objective}
- Ээлжит хичээлийн зорилго: {selected['Ээлжит хичээлийн зорилго']}
- Суралцахуйн үр дүн: {learning_outcome}

Төлөвлөлтийн загвар:
1. Хичээлийн үе шат (Хугацаа, Багшийн үйл ажиллагаа, Сурагчийн үйл ажиллагаа хүснэгттэй)
2. Үнэлгээний нэгж, шалгуур
3. Дасгал, бодлого, даалгавар
4. Дүгнэлт, гэрийн даалгавар
"""

        extra_text = st.text_area("Промтод нэмэх зүйлээ бичнэ үү", key="extra_prompt")
        prompt_full = default_prompt.strip() + "\n" + extra_text.strip() if extra_text else default_prompt.strip()

        if st.button("✍️ Gemini-р төлөвлөгөө гаргах"):
            with st.spinner("Төлөвлөгөө үүсгэж байна..."):
                response = model.generate_content(prompt_full)
                plan_text = response.text
                st.session_state['ai_plan'] = plan_text

        if st.session_state.get('ai_plan'):
            st.markdown("#### 📝 Gemini-гээс гаргасан төлөвлөгөө:")
            st.markdown(st.session_state["ai_plan"])

            # --- DOCX болгож хөрвүүлэх ---
            def generate_docx(text):
                doc = Document()
                doc.add_heading("Ээлжит хичээлийн төлөвлөгөө", 0)

                # Markdown хүснэгтүүдийг илрүүлэх
                tables = re.findall(r"((?:\|.*\|\n)+)", text)

                # Энгийн текст
                for part in text.split("\n"):
                    if "|" in part:
                        continue
                    elif part.strip():
                        doc.add_paragraph(part.strip())

                # Хүснэгтүүдийг хөрвүүлэх
                for table in tables:
                    lines = [line.strip() for line in table.strip().split("\n") if line.strip()]
                    headers = [h.strip() for h in lines[0].split("|") if h.strip()]
                    rows = []
                    for line in lines[2:]:
                        cols = [c.strip() for c in line.split("|") if c.strip()]
                        if len(cols) == len(headers):
                            rows.append(cols)

                    if headers and rows:
                        doc_table = doc.add_table(rows=1, cols=len(headers))
                        doc_table.style = "Table Grid"
                        hdr_cells = doc_table.rows[0].cells
                        for i, h in enumerate(headers):
                            hdr_cells[i].text = h
                        for row in rows:
                            row_cells = doc_table.add_row().cells
                            for i, val in enumerate(row):
                                row_cells[i].text = val
                        doc.add_paragraph("")

                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                return buf

            docx_file = generate_docx(st.session_state["ai_plan"])
            st.download_button(
                label="⬇️ DOCX татаж авах",
                data=docx_file,
                file_name="lesson_plan.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
